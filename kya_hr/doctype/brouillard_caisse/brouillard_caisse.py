import frappe
from frappe.model.document import Document
from frappe.utils import flt, today

from kya_hr.utils.approval_guards import block_self_approval


class BrouillardCaisse(Document):
    """Brouillard de Caisse — saisie quotidienne par la Caissière.

    Workflow: Caissière (saisit) → Comptable (vise) → DFC (valide).
    Totaux calculés + solde ligne-à-ligne côté serveur pour cohérence,
    même si le JS fait la même chose côté client.
    """

    def validate(self):
        block_self_approval(self)
        self._compute_line_soldes()
        self._compute_totals()
        self._auto_sign_caissiere()

    def _compute_line_soldes(self):
        solde = flt(self.solde_precedent or 0)
        for ligne in (self.lignes or []):
            solde += flt(ligne.entree or 0) - flt(ligne.sortie or 0)
            ligne.solde = solde

    def _compute_totals(self):
        self.total_entrees = sum(flt(l.entree or 0) for l in (self.lignes or []))
        self.total_sorties = sum(flt(l.sortie or 0) for l in (self.lignes or []))
        self.solde_final = flt(self.solde_precedent or 0) + flt(self.total_entrees) - flt(self.total_sorties)

    def _auto_sign_caissiere(self):
        if self.caissiere and not self.signataire_caissiere:
            emp_name = frappe.db.get_value("Employee", self.caissiere, "employee_name")
            self.signataire_caissiere = emp_name or self.caissiere
            self.date_signature_caissiere = today()


# ─── Point hebdomadaire DG ──────────────────────────────────────────────
def send_weekly_dg_summary():
    """Scheduler vendredi 17h: email récap court au DG + DGA, DAAF en copie.

    Le mail contient juste les chiffres clés et un bouton vers la page
    `/recap-brouillards?semaine=YYYY-Www` qui affiche le détail + exports
    PDF/DOCX.

    Déclenché par hooks.py -> scheduler_events.cron "0 17 * * 5".
    """
    from frappe.utils import add_days, getdate, formatdate

    today_d = getdate()
    weekday = today_d.weekday()  # lundi=0 ... vendredi=4
    start = add_days(today_d, -weekday)
    end = add_days(start, 4)
    iso_year, iso_week, _ = today_d.isocalendar()
    iso_tag = f"{iso_year}-W{iso_week:02d}"

    brouillards = frappe.get_all(
        "Brouillard Caisse",
        filters={
            "docstatus": 1,
            "date_brouillard": ["between", [start, end]],
        },
        fields=[
            "name", "total_entrees", "total_sorties", "solde_final",
        ],
    )

    if not brouillards:
        return

    recipients = _get_users_with_roles(["Directeur Général", "DGA"])
    cc = _get_users_with_roles(["DAAF"])
    if not recipients:
        return

    total_e = sum(flt(b.total_entrees or 0) for b in brouillards)
    total_s = sum(flt(b.total_sorties or 0) for b in brouillards)
    solde = total_e - total_s
    nb = len(brouillards)

    base = frappe.utils.get_url()
    page_url = f"{base}/recap-brouillards?semaine={iso_tag}"

    message = (
        "<div style='font-family:Arial,sans-serif;max-width:560px;margin:0 auto;'>"
        "<div style='background:linear-gradient(135deg,#009688,#00bcd4);padding:24px;"
        "border-radius:12px 12px 0 0;text-align:center;'>"
        f"<img src='{base}/assets/kya_hr/images/kya_logo.png' width='60' style='margin-bottom:8px;'>"
        "<h2 style='color:white;margin:0;'>Point Hebdomadaire — Caisse</h2>"
        f"<p style='color:white;margin:6px 0 0;font-size:14px;'>Semaine du {formatdate(start)} au {formatdate(end)}</p>"
        "</div>"
        "<div style='background:white;padding:24px;border:1px solid #e0e0e0;border-radius:0 0 12px 12px;'>"
        "<p>Bonjour,</p>"
        f"<p>Le récap des <b>{nb}</b> brouillards de caisse approuvés cette semaine est disponible.</p>"
        "<table style='width:100%;border-collapse:collapse;margin:14px 0;font-size:14px;'>"
        f"<tr><td style='padding:8px 12px;background:#f5f5f5;'>Total entrées</td>"
        f"<td style='padding:8px 12px;text-align:right;font-weight:700;color:#2e7d32;'>{total_e:,.0f} XOF</td></tr>"
        f"<tr><td style='padding:8px 12px;background:#f5f5f5;'>Total sorties</td>"
        f"<td style='padding:8px 12px;text-align:right;font-weight:700;color:#c62828;'>{total_s:,.0f} XOF</td></tr>"
        f"<tr><td style='padding:8px 12px;background:#e8f5e9;'><b>Solde net</b></td>"
        f"<td style='padding:8px 12px;text-align:right;font-weight:700;color:#1565c0;font-size:16px;'>{solde:,.0f} XOF</td></tr>"
        "</table>"
        "<div style='text-align:center;margin:24px 0;'>"
        f"<a href='{page_url}' style='display:inline-block;padding:14px 32px;"
        "background:#009688;color:white;text-decoration:none;border-radius:8px;"
        "font-weight:700;font-size:15px;'>📊 Consulter le récap détaillé</a>"
        "</div>"
        "<p style='font-size:12px;color:#666;text-align:center;'>"
        "Vous pouvez télécharger le récap en PDF ou Word depuis la page."
        "</p>"
        "</div></div>"
    ).replace(",", " ")

    frappe.sendmail(
        recipients=recipients,
        cc=cc,
        subject=f"[KYA] Point hebdo caisse — semaine du {formatdate(start)} ({nb} brouillards)",
        message=message,
        reference_doctype="Brouillard Caisse",
        now=False,
    )


def _get_users_with_roles(roles):
    if not roles:
        return []
    users = frappe.get_all(
        "Has Role",
        filters={"role": ["in", roles], "parenttype": "User"},
        fields=["parent"],
        distinct=True,
    )
    emails = []
    for u in users:
        email = frappe.db.get_value("User", u.parent, "email")
        if email and email not in emails and u.parent not in ("Guest", "Administrator"):
            emails.append(email)
    return emails


# ─── Export PDF / DOCX du récap hebdomadaire ────────────────────────────
@frappe.whitelist()
def export_recap(semaine=None, fmt="pdf"):
    """Exporte le récap hebdo des brouillards en PDF ou DOCX.

    Args:
        semaine: format ISO YYYY-Www (défaut: semaine courante).
        fmt: 'pdf' ou 'docx'.

    Renvoie le fichier en téléchargement direct.
    """
    from frappe.utils import add_days, getdate, formatdate
    from datetime import date

    allowed = {"Directeur Général", "DGA", "DAAF", "Comptable",
               "Responsable RH", "System Manager", "HR Manager"}
    if not allowed.intersection(set(frappe.get_roles(frappe.session.user))):
        frappe.throw("Accès refusé", frappe.PermissionError)

    today_d = getdate()
    if semaine and "-W" in semaine:
        try:
            year, wk = semaine.split("-W")
            year, wk = int(year), int(wk)
            jan4 = date(year, 1, 4)
            week1_monday = add_days(jan4, -jan4.weekday())
            start = add_days(week1_monday, (wk - 1) * 7)
        except Exception:
            start = add_days(today_d, -today_d.weekday())
    else:
        start = add_days(today_d, -today_d.weekday())
    end = add_days(start, 4)
    iso_year, iso_week, _w = getdate(start).isocalendar()

    brouillards = frappe.get_all(
        "Brouillard Caisse",
        filters={"docstatus": 1, "date_brouillard": ["between", [start, end]]},
        fields=["name", "date_brouillard", "caissiere_name",
                "total_entrees", "total_sorties", "solde_final", "statut"],
        order_by="date_brouillard asc",
    )
    total_e = sum(flt(b.total_entrees or 0) for b in brouillards)
    total_s = sum(flt(b.total_sorties or 0) for b in brouillards)
    solde = total_e - total_s

    filename_base = f"recap-brouillards-{iso_year}-W{iso_week:02d}"

    if fmt == "docx":
        return _export_recap_docx(brouillards, start, end, iso_week, iso_year,
                                  total_e, total_s, solde, filename_base)
    return _export_recap_pdf(brouillards, start, end, iso_week, iso_year,
                             total_e, total_s, solde, filename_base)


def _export_recap_pdf(brouillards, start, end, iso_week, iso_year,
                      total_e, total_s, solde, filename_base):
    from frappe.utils import formatdate
    from frappe.utils.pdf import get_pdf

    rows = "".join(
        "<tr>"
        f"<td>{formatdate(b.date_brouillard)}</td>"
        f"<td>{b.name}</td>"
        f"<td>{b.caissiere_name or ''}</td>"
        f"<td style='text-align:right;'>{flt(b.total_entrees or 0):,.0f}</td>"
        f"<td style='text-align:right;'>{flt(b.total_sorties or 0):,.0f}</td>"
        f"<td style='text-align:right;'><b>{flt(b.solde_final or 0):,.0f}</b></td>"
        f"<td>{b.statut or ''}</td>"
        "</tr>"
        for b in brouillards
    )
    html = f"""
    <html><head><meta charset='utf-8'><style>
    body {{ font-family: Arial, sans-serif; font-size: 12px; padding: 24px; }}
    h1 {{ color: #009688; margin-bottom: 4px; }}
    .meta {{ color: #666; margin-bottom: 18px; }}
    table {{ width: 100%; border-collapse: collapse; }}
    th {{ background: #009688; color: white; padding: 8px; text-align: left; }}
    td {{ padding: 6px 8px; border-bottom: 1px solid #eee; }}
    tfoot td {{ background: #f1f8e9; font-weight: 700; }}
    .kpi {{ display: inline-block; margin: 4px 12px 14px 0; padding: 8px 14px; background: #f5f5f5; border-radius: 6px; }}
    </style></head><body>
    <h1>Récap Hebdomadaire — Caisse</h1>
    <div class='meta'>Semaine {iso_week} / {iso_year} — du {formatdate(start)} au {formatdate(end)}</div>
    <div>
        <span class='kpi'>Entrées : <b>{total_e:,.0f} XOF</b></span>
        <span class='kpi'>Sorties : <b>{total_s:,.0f} XOF</b></span>
        <span class='kpi'>Solde : <b>{solde:,.0f} XOF</b></span>
        <span class='kpi'>Brouillards : <b>{len(brouillards)}</b></span>
    </div>
    <table><thead><tr>
    <th>Date</th><th>Réf.</th><th>Caissier(ère)</th>
    <th style='text-align:right;'>Entrées</th><th style='text-align:right;'>Sorties</th>
    <th style='text-align:right;'>Solde Final</th><th>Statut</th>
    </tr></thead>
    <tbody>{rows or '<tr><td colspan=7 style="text-align:center;padding:30px;">Aucun brouillard</td></tr>'}</tbody>
    <tfoot><tr>
    <td colspan='3'>Totaux semaine</td>
    <td style='text-align:right;'>{total_e:,.0f}</td>
    <td style='text-align:right;'>{total_s:,.0f}</td>
    <td style='text-align:right;'>{solde:,.0f}</td>
    <td></td></tr></tfoot>
    </table>
    </body></html>
    """.replace(",", " ")

    pdf_bytes = get_pdf(html)
    frappe.local.response.filename = f"{filename_base}.pdf"
    frappe.local.response.filecontent = pdf_bytes
    frappe.local.response.type = "download"


def _export_recap_docx(brouillards, start, end, iso_week, iso_year,
                       total_e, total_s, solde, filename_base):
    from frappe.utils import formatdate
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor
    except ImportError:
        frappe.throw("python-docx n'est pas installé. Lancez : pip install python-docx")

    doc = DocxDocument()
    title = doc.add_heading("Récap Hebdomadaire — Caisse", level=1)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x00, 0x96, 0x88)
    doc.add_paragraph(
        f"Semaine {iso_week} / {iso_year} — du {formatdate(start)} au {formatdate(end)}"
    )

    p = doc.add_paragraph()
    p.add_run(f"Entrées : {total_e:,.0f} XOF    ".replace(",", " ")).bold = True
    p.add_run(f"Sorties : {total_s:,.0f} XOF    ".replace(",", " ")).bold = True
    p.add_run(f"Solde : {solde:,.0f} XOF    ".replace(",", " ")).bold = True
    p.add_run(f"Brouillards : {len(brouillards)}").bold = True

    table = doc.add_table(rows=1, cols=7)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Date", "Réf.", "Caissier(ère)", "Entrées",
                           "Sorties", "Solde Final", "Statut"]):
        hdr[i].text = h

    for b in brouillards:
        row = table.add_row().cells
        row[0].text = formatdate(b.date_brouillard)
        row[1].text = b.name or ""
        row[2].text = b.caissiere_name or ""
        row[3].text = f"{flt(b.total_entrees or 0):,.0f}".replace(",", " ")
        row[4].text = f"{flt(b.total_sorties or 0):,.0f}".replace(",", " ")
        row[5].text = f"{flt(b.solde_final or 0):,.0f}".replace(",", " ")
        row[6].text = b.statut or ""

    # Ligne totaux
    foot = table.add_row().cells
    foot[0].text = "TOTAL"
    foot[3].text = f"{total_e:,.0f}".replace(",", " ")
    foot[4].text = f"{total_s:,.0f}".replace(",", " ")
    foot[5].text = f"{solde:,.0f}".replace(",", " ")
    for c in foot:
        for para in c.paragraphs:
            for run in para.runs:
                run.bold = True

    import io
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    frappe.local.response.filename = f"{filename_base}.docx"
    frappe.local.response.filecontent = buf.getvalue()
    frappe.local.response.type = "download"


